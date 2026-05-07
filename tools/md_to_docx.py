"""將 reports/<rid>_run/*.md 整併轉成單份 docx。

針對 md_reporter.py 的輸出格式做 line-based parser，支援：
- # / ## / ### 標題
- | 表格（含 separator 行）
- ```...``` code block
- ![alt](path) 圖片（path 相對 md 解析；用 BytesIO 顯式嵌入）
- - 列表
- 一般段落

整併規則：`_summary.md` 在最前，其餘工項報告依檔名排序，之間插 page break。
輸出單檔：`reports/<rid>_run/docx/<rid>_run_測試報告.docx`

不支援巢狀結構或一般化 markdown；reporter 不會產出複雜 markdown，足夠交付用。

用法：
    python tools/md_to_docx.py                 # 轉最新一筆 run
    python tools/md_to_docx.py 20260506_2030   # 轉指定 run
"""
from __future__ import annotations
import argparse
import re
import sys
from io import BytesIO
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

try:
    from docx import Document
    from docx.shared import Inches, Pt
except ImportError:
    print(
        "缺少 python-docx 依賴。請執行：pip install python-docx==1.1.2",
        file=sys.stderr,
    )
    raise

REPORTS_DIR = PROJECT_ROOT / "reports"

H_RE = re.compile(r"^(#{1,3})\s+(.*)$")
TABLE_RE = re.compile(r"^\|(.*)\|\s*$")
TABLE_SEP_RE = re.compile(r"^\|[\s\-:|]+\|\s*$")
# 用 greedy `.+`：alt 與 path 都可能含 `[chromium]`（pytest-playwright 標籤），
# 不能用 `[^\]]*`/`[^)]+`——會在第一個 `]` 就停而對不上後面 `](`
IMG_RE = re.compile(r"^!\[(.+)\]\((.+)\)\s*$")
LIST_RE = re.compile(r"^- (.*)$")
FENCE_RE = re.compile(r"^```")
LINK_INLINE_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")
BOLD_INLINE_RE = re.compile(r"\*\*(.+?)\*\*")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")


def _strip_inline_md(text: str) -> str:
    text = LINK_INLINE_RE.sub(r"\1", text)
    text = BOLD_INLINE_RE.sub(r"\1", text)
    text = INLINE_CODE_RE.sub(r"\1", text)
    return text


def _embed_picture(doc, img_path: Path) -> None:
    """讀檔成 BytesIO 再丟給 add_picture，確保是嵌入而非任何形式的連結。"""
    with open(img_path, "rb") as f:
        buf = BytesIO(f.read())
    doc.add_picture(buf, width=Inches(6))


def _append_md_to_doc(doc, md_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    base_dir = md_path.parent

    in_code = False
    code_buf: list[str] = []
    i = 0

    while i < len(lines):
        line = lines[i]

        if FENCE_RE.match(line):
            if in_code:
                if code_buf:
                    p = doc.add_paragraph()
                    run = p.add_run("\n".join(code_buf))
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        m = H_RE.match(line)
        if m:
            level = len(m.group(1))
            doc.add_heading(_strip_inline_md(m.group(2)), level=level)
            i += 1
            continue

        m = IMG_RE.match(line)
        if m:
            alt = m.group(1)
            rel = m.group(2)
            img_path = (base_dir / rel).resolve()
            if img_path.exists():
                try:
                    _embed_picture(doc, img_path)
                    if alt:
                        cap = doc.add_paragraph(alt)
                        if cap.runs:
                            cap.runs[0].italic = True
                except Exception as e:
                    doc.add_paragraph(f"[圖片載入失敗：{img_path.name} — {e}]")
            else:
                doc.add_paragraph(f"[找不到圖片：{rel}]")
            i += 1
            continue

        if TABLE_RE.match(line):
            tbl_lines = []
            while i < len(lines) and TABLE_RE.match(lines[i]):
                tbl_lines.append(lines[i])
                i += 1
            _emit_table(doc, tbl_lines)
            continue

        m = LIST_RE.match(line)
        if m:
            doc.add_paragraph(_strip_inline_md(m.group(1)), style="List Bullet")
            i += 1
            continue

        if line.strip():
            doc.add_paragraph(_strip_inline_md(line))
        else:
            doc.add_paragraph("")
        i += 1


def _emit_table(doc, table_lines: list[str]) -> None:
    rows = []
    for ln in table_lines:
        if TABLE_SEP_RE.match(ln):
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=cols)
    try:
        tbl.style = "Light Grid Accent 1"
    except KeyError:
        pass
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            tbl.rows[r_idx].cells[c_idx].text = _strip_inline_md(cell)


def _ordered_md_files(run_dir: Path) -> list[Path]:
    summary = run_dir / "_summary.md"
    rest = sorted(p for p in run_dir.glob("*.md") if p.name != "_summary.md")
    return ([summary] if summary.exists() else []) + rest


def convert_run(run_id: str) -> Path:
    if run_id.endswith("_run"):
        run_dir = REPORTS_DIR / run_id
    else:
        run_dir = REPORTS_DIR / f"{run_id}_run"
    if not run_dir.is_dir():
        raise SystemExit(f"找不到 run 目錄：{run_dir}")

    md_files = _ordered_md_files(run_dir)
    if not md_files:
        raise SystemExit(f"{run_dir} 內無 md 檔可轉")

    doc = Document()
    for idx, md in enumerate(md_files):
        if idx > 0:
            doc.add_page_break()
        print(f"  → 併入 {md.name}")
        _append_md_to_doc(doc, md)

    out_dir = run_dir / "docx"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / f"{run_dir.name}_測試報告.docx"
    doc.save(str(out_path))
    print(f"\n完成：{out_path.relative_to(PROJECT_ROOT)}")
    return out_path


def latest_run_id() -> str:
    runs = [p.name for p in REPORTS_DIR.iterdir() if p.is_dir() and p.name.endswith("_run")]
    if not runs:
        raise SystemExit("reports/ 內找不到任何 _run 目錄")
    return sorted(runs)[-1]


def main() -> int:
    p = argparse.ArgumentParser(description="md_reporter 報告整併成單份 docx")
    p.add_argument(
        "run_id",
        nargs="?",
        help="run 目錄名（如 20260506_2030 或 20260506_2030_run）；省略則用最新一筆",
    )
    args = p.parse_args()

    run_id = args.run_id or latest_run_id()
    convert_run(run_id)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
